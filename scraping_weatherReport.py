# importing all the necessary packages
import urllib.request # This module helps to define functions and classes to open URLs (mostly http)
import docx  # It creates, reads and wrires the Microsoft office Word docx files
from docx.shared import Inches # image adujstment
from datetime import datetime # To retrive the date and time
import json # transmit and receive data between a server and web application in JSON format.
import time

# creating a document:
doc = docx.Document()
date_time = datetime.now().strftime("%d %b %Y | %I:%M:%S %p")
# heading of the docx
doc.add_heading('⛅🌦️🏖️WEATHER FORECASTING DETAILS🏖️🌥️🌤️-----------------------------' + date_time, 0)
# adding pictures
doc.add_picture("img.jpg", width=Inches(3.5))
# write a paragraph:
weather = doc.add_paragraph()
APIKey = "714c2da2279a72ec8bb713798409d352"



# defining the function to scrap the weather report from the web using api key and url
def web_scrap():
    location = input("Enter location : ")
    url = "https://api.openweathermap.org/data/2.5/weather?q=%s&APPID=%s" % (location, APIKey)  # Service url
    req = urllib.request.urlopen(url).read().decode()#passing a dictionary parameter to urllib.urlopen
    data = json.loads(req)# returns the dictionary
    riseTime = time.strftime("%D %H:%M",time.localtime(int(data['sys']['sunrise']))).split()  # formatting the rise time
    setTime = time.strftime("%D %H:%M", time.localtime(int(data['sys']['sunset']))).split()  # set Time
    # Printing the generated weather report on document
    weather.add_run(
        "\nLocation: %s\nWeather: %s\nTemperature: %s°F | %s°C\nWind Speed: %s MPH\nWind Degree: %s°Degrees\nSunrise: %s\nSunset: %s\nLatitude: %s\nLongitude: %s" % (
            data["name"], data["weather"][0]["main"], round((data["main"]["temp"] - 273.15) * 9 / 5 + 32, 2),
            round(data["main"]["temp"] - 273.15, 2), data["wind"]["speed"], data["wind"]["deg"], riseTime[1],
            setTime[1], data['coord']['lat'], data['coord']['lon'])).arial = True
    weather.add_run("\n----------------------------------------------------------------------------------\n")


